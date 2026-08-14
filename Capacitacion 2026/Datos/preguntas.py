import pandas as pd
from docx import Document

# Cargar el archivo Excel
df = pd.read_excel("preguntas.xlsx")

# Crear documento Word
doc = Document()

for idx, row in df.iterrows():
    pregunta = str(row["Pregunta"]).strip()
    opciones = [row.get("A"), row.get("B"), row.get("C"), row.get("D")]

    # Agregar la pregunta con el prefijo "Pregunta:"
    doc.add_paragraph(f"Pregunta: {pregunta}")

    # Agregar las opciones con guion
    for opcion in opciones:
        if pd.notna(opcion):
            doc.add_paragraph(f"- {str(opcion).strip()}")

    # Línea en blanco para separar preguntas
    doc.add_paragraph("")

# Guardar el archivo Word
doc.save("preguntas_forms.docx")

print("Archivo 'preguntas_forms.docx' generado correctamente.")

